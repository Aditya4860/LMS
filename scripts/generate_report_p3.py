from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

document = Document('Library_Management_System_Project_Report.docx')
fig_count = 1
tab_count = 2

def add_chapter_heading(doc, text):
    doc.add_page_break()
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_subheading(doc, text):
    doc.add_heading(text, level=2)

def add_subsubheading(doc, text):
    doc.add_heading(text, level=3)

def add_text(doc, text):
    doc.add_paragraph(text)

def add_image_with_caption(doc, img_path, caption):
    global fig_count
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        try:
            r.add_picture(img_path, width=Inches(6.0))
        except Exception as e:
            print(f"Error adding image {img_path}: {e}")
        cap_p = doc.add_paragraph(f"Figure {fig_count}: {caption}")
        cap_p.style = 'Caption'
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fig_count += 1
    else:
        doc.add_paragraph(f"[Image placeholder for {caption}]")

def add_table_with_caption(doc, caption, headers, rows):
    global tab_count
    cap_p = doc.add_paragraph(f"Table {tab_count}: {caption}")
    cap_p.style = 'Caption'
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    doc.add_paragraph("")
    tab_count += 1

# Chapter 3
add_chapter_heading(document, "CHAPTER 3: SYSTEM ANALYSIS")
add_subheading(document, "3.1 Functional Requirements")
add_text(document, "1. The system must allow users to register and login.\n2. The system must allow librarians to add, edit, and delete books.\n3. Students must be able to search for books and view availability.\n4. Librarians must be able to issue books and process returns.\n5. The system must automatically calculate fines for late returns.")
add_subheading(document, "3.2 Non Functional Requirements")
add_text(document, "1. Security: Passwords must be hashed using ASP.NET Identity.\n2. Performance: Queries must use pagination to handle large datasets.\n3. Reliability: The system should have 99% uptime.\n4. Usability: The UI must be responsive and intuitive.")
add_subheading(document, "3.3 Software Requirements")
add_table_with_caption(document, "Software Requirements", 
    ["Component", "Specification"],
    [["Operating System", "Windows 10/11 or Linux"],
     ["Framework", ".NET 8.0 / ASP.NET Core MVC"],
     ["Database", "Microsoft SQL Server"],
     ["Frontend", "HTML5, CSS3, Bootstrap 5, JavaScript"],
     ["IDE", "Visual Studio 2022 / VS Code"]])
add_subheading(document, "3.4 Hardware Requirements")
add_table_with_caption(document, "Hardware Requirements", 
    ["Component", "Minimum Requirement"],
    [["Processor", "Intel Core i3 or equivalent"],
     ["RAM", "4 GB (8 GB Recommended)"],
     ["Storage", "500 MB Free Space"],
     ["Network", "Internet Connection"]])
add_subheading(document, "3.5 Feasibility Study")
add_text(document, "The project is technically feasible as it utilizes well-established technologies. It is economically feasible due to the use of open-source and standard tools, avoiding extreme licensing costs. Operationally, it satisfies the requirements of a modern library.")

add_subheading(document, "3.6 DFD")
add_image_with_caption(document, "Generated_Diagrams/dfd_level_0.png", "DFD Level 0")
add_image_with_caption(document, "Generated_Diagrams/dfd_level_1.png", "DFD Level 1")

add_subheading(document, "3.7 Use Case Diagram")
add_image_with_caption(document, "Generated_Diagrams/use_case_diagram.png", "Use Case Diagram")

add_subheading(document, "3.8 Activity Diagram")
add_image_with_caption(document, "Generated_Diagrams/activity_diagram.png", "Activity Diagram")

add_subheading(document, "3.9 Sequence Diagram")
add_image_with_caption(document, "Generated_Diagrams/sequence_diagram.png", "Sequence Diagram")

add_subheading(document, "3.10 Class Diagram")
add_image_with_caption(document, "Generated_Diagrams/class_diagram.png", "Class Diagram")

add_subheading(document, "3.11 Architecture Diagram")
add_image_with_caption(document, "Generated_Diagrams/system_architecture.png", "System Architecture")

add_subheading(document, "3.12 Database Design")
add_text(document, "The database is designed using Entity Framework Core Code-First approach. Below is the Entity-Relationship Diagram depicting the structure.")
add_image_with_caption(document, "Generated_Diagrams/er_diagram.png", "ER Diagram")
add_image_with_caption(document, "Generated_Diagrams/database_relationship_diagram.png", "Database Relationship Diagram")

add_subheading(document, "3.13 Database Schema & Tables")
add_table_with_caption(document, "Books Table Schema", 
    ["Column Name", "Data Type", "Constraints"],
    [["Id", "int", "Primary Key, Identity"],
     ["Title", "nvarchar(max)", "Required"],
     ["Author", "nvarchar(max)", "Required"],
     ["ISBN", "nvarchar(100)", "Required"],
     ["TotalCopies", "int", "Required"]])

add_table_with_caption(document, "BorrowRecords Table Schema", 
    ["Column Name", "Data Type", "Constraints"],
    [["Id", "int", "Primary Key, Identity"],
     ["BookId", "int", "Foreign Key (Books)"],
     ["UserId", "nvarchar(450)", "Foreign Key (AspNetUsers)"],
     ["BorrowDate", "datetime2", "Required"],
     ["ReturnDate", "datetime2", "Nullable"]])

document.save('Library_Management_System_Project_Report.docx')
print("Part 3 completed.")
