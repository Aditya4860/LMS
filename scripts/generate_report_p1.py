import os
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_toc(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_list_of_figures(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\h \\z \\c "Figure"'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def add_list_of_tables(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\h \\z \\c "Table"'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

document = Document()

# Page setup
for section in document.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17) # Left margin slightly larger for binding
    section.right_margin = Cm(2.54)

# Set Default Font
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# --- Cover Page ---
document.add_heading('LIBRARY MANAGEMENT SYSTEM', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
p = document.add_paragraph('A Project Report\nSubmitted in partial fulfillment of the requirements for the degree of\nBACHELOR OF TECHNOLOGY\nin\nCOMPUTER SCIENCE AND ENGINEERING\n')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
document.add_page_break()

# --- Certificate ---
document.add_heading('CERTIFICATE', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
document.add_paragraph('This is to certify that the project entitled "LIBRARY MANAGEMENT SYSTEM" is a bonafide record of the work done, submitted in partial fulfillment of the requirements for the award of Bachelor of Technology in Computer Science and Engineering.')
document.add_page_break()

# --- Acknowledgement ---
document.add_heading('ACKNOWLEDGEMENT', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
document.add_paragraph('I would like to express my sincere gratitude to my supervisor, the department, and the university for their support and guidance throughout this project.')
document.add_page_break()

# --- Abstract ---
document.add_heading('ABSTRACT', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
document.add_paragraph('The Library Management System (LMS) is a comprehensive web-based application designed to manage the core operations of a library. The system digitizes processes such as book tracking, member registration, borrowing, and returning mechanisms. Developed using ASP.NET Core MVC, Entity Framework Core, and SQL Server, it provides a robust, scalable, and user-friendly interface. Role-based access control (RBAC) ensures distinct experiences for Administrators, Librarians, and Students. Features include real-time availability tracking, penalty calculation, advanced search, and pagination. Thorough unit testing using xUnit ensures system reliability. The implementation of this system minimizes manual errors and enhances operational efficiency.')
document.add_page_break()

# --- TOC, LOF, LOT ---
document.add_heading('TABLE OF CONTENTS', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_toc(document)
document.add_page_break()

document.add_heading('LIST OF FIGURES', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_list_of_figures(document)
document.add_page_break()

document.add_heading('LIST OF TABLES', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_list_of_tables(document)
document.add_page_break()

document.add_heading('ABBREVIATIONS', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
abbr_table = document.add_table(rows=1, cols=2)
abbr_table.style = 'Table Grid'
hdr_cells = abbr_table.rows[0].cells
hdr_cells[0].text = 'Abbreviation'
hdr_cells[1].text = 'Description'
abbrs = [('LMS', 'Library Management System'), ('MVC', 'Model View Controller'), ('EF', 'Entity Framework'), ('UI', 'User Interface'), ('DFD', 'Data Flow Diagram'), ('ER', 'Entity Relationship')]
for a, d in abbrs:
    row_cells = abbr_table.add_row().cells
    row_cells[0].text = a
    row_cells[1].text = d
document.add_page_break()

document.save('Library_Management_System_Project_Report.docx')
print("Document partially created successfully.")
