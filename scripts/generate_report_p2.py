from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

document = Document('Library_Management_System_Project_Report.docx')
fig_count = 1
tab_count = 1

def add_chapter_heading(doc, text):
    doc.add_page_break()
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_subheading(doc, text):
    doc.add_heading(text, level=2)

def add_subsubheading(doc, text):
    doc.add_heading(text, level=3)

def add_text(doc, text):
    doc.add_paragraph(text)

def add_image_with_caption(doc, img_path, caption):
    global fig_count
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        try:
            r.add_picture(img_path, width=Inches(6.0))
        except Exception as e:
            print(f"Error adding image {img_path}: {e}")
        cap_p = doc.add_paragraph(f"Figure {fig_count}: {caption}")
        cap_p.style = 'Caption'
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fig_count += 1
    else:
        doc.add_paragraph(f"[Image placeholder for {caption}]")

def add_table_with_caption(doc, caption, headers, rows):
    global tab_count
    cap_p = doc.add_paragraph(f"Table {tab_count}: {caption}")
    cap_p.style = 'Caption'
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    for row in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    doc.add_paragraph("")
    tab_count += 1

# Chapter 1
add_chapter_heading(document, "CHAPTER 1: INTRODUCTION")
add_subheading(document, "1.1 Project Overview")
add_text(document, "The Library Management System is a modern digital solution aimed at automating the core day-to-day functions of a library. The software enables members, librarians, and administrators to seamlessly engage with the library's catalog, borrowing procedures, and administrative duties.")
add_subheading(document, "1.2 Problem Statement")
add_text(document, "Traditional libraries often rely on manual ledgers or outdated systems which lead to inefficiencies, high error rates, slow search processes, and difficulties in tracking fines or inventory. This project addresses the need for a scalable, automated system.")
add_subheading(document, "1.3 Objectives")
add_text(document, "The primary objectives are to automate library operations, provide real-time book availability, implement role-based access control, enforce validation for data integrity, and provide pagination and advanced search capabilities.")
add_subheading(document, "1.4 Need of the System")
add_text(document, "A computerized LMS is necessary to eliminate manual labor, streamline issuing and returning of resources, automatically calculate penalties, and provide detailed reporting mechanisms for management.")
add_subheading(document, "1.5 Scope")
add_text(document, "The system encompasses managing books, magazines, newspapers, students, and librarians. It handles authentication, borrowing records, returns, search, pagination, and administrative dashboards.")
add_subheading(document, "1.6 Existing System")
add_text(document, "The existing system relies on paper-based records, requiring users to manually search card catalogs and librarians to record transactions in physical registers.")
add_subheading(document, "1.7 Limitations")
add_text(document, "Physical registers are prone to damage, search is incredibly slow, calculating fines is tedious, and providing timely reports is nearly impossible.")
add_subheading(document, "1.8 Proposed System")
add_text(document, "The proposed system is a web-based application built with ASP.NET Core MVC, Entity Framework Core, and SQL Server. It digitizes the entire library ecosystem.")
add_subheading(document, "1.9 Advantages")
add_text(document, "It offers fast access, 24/7 availability, accuracy in fine calculations, a user-friendly UI, secure data management, and automated report generation.")

# Chapter 2
add_chapter_heading(document, "CHAPTER 2: LITERATURE SURVEY")
add_subheading(document, "2.1 Related Work")
add_text(document, "Several systems have been developed to automate libraries. Early systems relied on desktop applications built with MS Access or Visual Basic. Modern architectures have migrated towards web-based cloud environments ensuring better scalability and multi-platform access.")
add_subheading(document, "2.2 Comparison Table")
add_table_with_caption(document, "System Comparison", 
    ["Feature", "Traditional System", "Desktop LMS", "Proposed Web-based LMS"],
    [
        ["Accessibility", "Physical presence required", "Local network only", "Anywhere via web browser"],
        ["Search Speed", "Slow", "Medium", "Fast (Pagination and Indexing)"],
        ["Data Security", "Low (Paper)", "Medium", "High (SQL Server, ASP.NET Identity)"],
        ["Fine Calculation", "Manual", "Automated", "Automated with real-time updates"]
    ])

document.save('Library_Management_System_Project_Report.docx')
print("Part 2 completed.")
